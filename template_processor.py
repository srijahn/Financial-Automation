"""
Template Processing Module - Handles DOCX and TXT template population
"""

import re
import logging
from pathlib import Path
from typing import Dict, Any, List
import pandas as pd

logger = logging.getLogger(__name__)

class EnhancedTemplateProcessor:
    """Enhanced template processor with DOCX support"""
    
    def __init__(self):
        self.placeholder_pattern = r'\{\{([^}]+)\}\}'
        try:
            from docx import Document
            self._has_docx = True
            self.Document = Document
        except ImportError:
            self._has_docx = False
            logger.warning("DOCX template support not available")
    
    def identify_fields(self, template_path: str) -> List[str]:
        """Identify template fields"""
        template_path = Path(template_path)
        
        if template_path.suffix.lower() == '.docx' and self._has_docx:
            return self._identify_docx_fields(template_path)
        else:
            return self._identify_txt_fields(template_path)
    
    def _identify_docx_fields(self, template_path: Path) -> List[str]:
        """Identify fields in DOCX template"""
        try:
            doc = self.Document(template_path)
            fields = set()
            
            # Check paragraphs
            for paragraph in doc.paragraphs:
                matches = re.findall(self.placeholder_pattern, paragraph.text)
                fields.update(matches)
            
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(self.placeholder_pattern, cell.text)
                        fields.update(matches)
            
            logger.info(f"Found {len(fields)} template fields in {template_path.name}")
            return list(fields)
            
        except Exception as e:
            logger.error(f"Error identifying DOCX fields: {e}")
            raise
    
    def _identify_txt_fields(self, template_path: Path) -> List[str]:
        """Identify fields in text template"""
        try:
            with open(template_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            fields = re.findall(self.placeholder_pattern, content)
            logger.info(f"Found {len(fields)} template fields in {template_path.name}")
            return list(set(fields))
            
        except Exception as e:
            logger.error(f"Error identifying text fields: {e}")
            raise
    
    def populate_template(self, template_path: str, data: Dict[str, Any], output_path: str):
        """Populate template with data and save as both Excel and DOCX"""
        template_path = Path(template_path)
        
        try:
            # Always create Excel output for data analysis
            df = pd.DataFrame(list(data.items()), columns=['Field Name', 'Value'])
            excel_output_path = Path(output_path).with_suffix('.xlsx')
            df.to_excel(excel_output_path, index=False)
            logger.info(f"Excel report saved: {excel_output_path}")
            
            # If template is DOCX, create populated DOCX document
            if template_path.suffix.lower() == '.docx' and self._has_docx:
                docx_output_path = Path(output_path).with_suffix('.docx')
                self._populate_docx_template(template_path, data, docx_output_path)
                logger.info(f"DOCX document saved: {docx_output_path}")
            
            # If template is TXT, create populated TXT document
            elif template_path.suffix.lower() == '.txt':
                txt_output_path = Path(output_path).with_suffix('.txt')
                self._populate_txt_template(template_path, data, txt_output_path)
                logger.info(f"TXT document saved: {txt_output_path}")
                
        except Exception as e:
            logger.error(f"Error populating template: {e}")
            raise

    def _populate_docx_template(self, template_path: Path, data: Dict[str, Any], output_path: Path):
        """Populate DOCX template with data"""
        try:
            # Load the template
            doc = self.Document(template_path)
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_text(paragraph, data)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_text(paragraph, data)
            
            # Replace placeholders in headers and footers
            for section in doc.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        self._replace_placeholders_in_text(paragraph, data)
                
                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        self._replace_placeholders_in_text(paragraph, data)
            
            # Save the populated document
            doc.save(output_path)
            
        except Exception as e:
            logger.error(f"Error populating DOCX template: {e}")
            raise

    def _populate_txt_template(self, template_path: Path, data: Dict[str, Any], output_path: Path):
        """Populate TXT template with data"""
        try:
            # Read template content
            with open(template_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Replace all placeholders
            for field_name, value in data.items():
                placeholder = f"{{{{{field_name}}}}}"
                content = content.replace(placeholder, str(value) if value else "")
            
            # Write populated content
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(content)
                
        except Exception as e:
            logger.error(f"Error populating TXT template: {e}")
            raise

    def _replace_placeholders_in_text(self, paragraph, data: Dict[str, Any]):
        """Replace placeholders in a paragraph while preserving formatting"""
        try:
            # Check if paragraph contains any placeholders
            if '{{' in paragraph.text and '}}' in paragraph.text:
                # Replace each placeholder
                for field_name, value in data.items():
                    placeholder = f"{{{{{field_name}}}}}"
                    if placeholder in paragraph.text:
                        # Replace placeholder with actual value
                        paragraph.text = paragraph.text.replace(
                            placeholder, 
                            str(value) if value else ""
                        )
        except Exception as e:
            logger.warning(f"Error replacing placeholders in paragraph: {e}")